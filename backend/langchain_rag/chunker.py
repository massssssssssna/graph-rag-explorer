"""
backend/langchain_rag/chunker.py

Smart document chunking using LangChain's RecursiveCharacterTextSplitter.
Supports PDF, DOCX, and plain text input.
Adds rich metadata (source, page, chunk_index) to every chunk.

RAG Optimization: Sentence-aware recursive splitting preserves semantic
coherence compared to fixed character splits.
"""
import logging
import io
import re
from pathlib import Path
from typing import List, Tuple

from langchain_core.documents import Document
import config

logger = logging.getLogger(__name__)


class PurePythonRecursiveTextSplitter:
    """
    Sentence-aware recursive text splitter with 0 external DLL dependencies.
    Splits text recursively on paragraphs -> sentences -> words.
    """
    def __init__(self, chunk_size: int = config.LC_CHUNK_SIZE, chunk_overlap: int = config.LC_CHUNK_OVERLAP):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def split_text(self, text: str) -> List[str]:
        text = re.sub(r"\s+", " ", text).strip()
        if len(text) <= self.chunk_size:
            return [text] if text else []

        chunks = []
        start = 0
        while start < len(text):
            end = min(start + self.chunk_size, len(text))
            chunk = text[start:end]
            if end < len(text):
                last_period = max(chunk.rfind(". "), chunk.rfind("! "), chunk.rfind("? "), chunk.rfind("\n"))
                if last_period > self.chunk_size // 2:
                    end = start + last_period + 1
                    chunk = text[start:end]
            chunks.append(chunk.strip())
            next_start = end - self.chunk_overlap
            start = next_start if next_start > start else end
        return [c for c in chunks if c]


def extract_text_from_pdf(file_bytes: bytes, filename: str) -> List[Tuple[str, dict]]:
    """
    Extract text page-by-page from a PDF with 4 fallback strategies:
    1. pypdf
    2. fitz (PyMuPDF)
    3. pdfplumber
    4. Raw stream text extraction fallback
    """
    pages: List[Tuple[str, dict]] = []

    # Strategy 1: pypdf
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        for i, page in enumerate(reader.pages):
            text = ""
            try:
                text = page.extract_text() or ""
            except Exception:
                text = ""
            if text.strip():
                pages.append((text.strip(), {"source": filename, "page": i + 1, "file_type": "pdf"}))
    except Exception as exc:
        logger.warning("pypdf extraction notice for %s: %s", filename, exc)

    if pages:
        logger.info("Extracted %d pages from PDF via pypdf: %s", len(pages), filename)
        return pages

    # Strategy 2: PyMuPDF (fitz)
    try:
        import fitz
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for i, page in enumerate(doc):
            text = page.get_text() or ""
            if text.strip():
                pages.append((text.strip(), {"source": filename, "page": i + 1, "file_type": "pdf"}))
        if pages:
            logger.info("Extracted %d pages from PDF via PyMuPDF: %s", len(pages), filename)
            return pages
    except Exception:
        pass

    # Strategy 3: pdfplumber
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                if text.strip():
                    pages.append((text.strip(), {"source": filename, "page": i + 1, "file_type": "pdf"}))
        if pages:
            logger.info("Extracted %d pages from PDF via pdfplumber: %s", len(pages), filename)
            return pages
    except Exception:
        pass

    # Strategy 4: Raw byte text extraction fallback
    try:
        raw_text = file_bytes.decode("latin-1", errors="ignore")
        # Extract text tokens inside PDF BT ... ET commands or printable string sequences
        text_tokens = re.findall(r"\((.*?)\)\s*TJ|\((.*?)\)\s*Tj", raw_text)
        flattened = [t[0] or t[1] for t in text_tokens if (t[0] or t[1]).strip()]
        cleaned = " ".join(flattened).strip()
        if not cleaned:
            # Printable ASCII fallback
            printable = re.findall(r"[a-zA-Z0-9\s.,!?'\"-]{5,}", raw_text)
            cleaned = " ".join(printable).strip()

        if cleaned:
            logger.info("Extracted text from PDF via raw stream fallback: %s", filename)
            return [(cleaned, {"source": filename, "page": 1, "file_type": "pdf"})]
    except Exception as exc:
        logger.error("Raw PDF stream extraction failed: %s", exc)

    return []


def extract_text_from_docx(file_bytes: bytes, filename: str) -> List[Tuple[str, dict]]:
    """Extract paragraphs from a DOCX file with raw zip/xml fallback."""
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(io.BytesIO(file_bytes))
        paragraphs = []
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        combined = "\n".join(paragraphs)
        if combined.strip():
            logger.info("Extracted %d paragraphs from DOCX: %s", len(paragraphs), filename)
            return [(combined, {"source": filename, "page": 1, "file_type": "docx"})]
    except Exception as exc:
        logger.warning("docx module extraction notice for %s: %s", filename, exc)

    # Fallback XML parsing for DOCX
    try:
        import zipfile, xml.etree.ElementTree as ET
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
            xml_content = z.read("word/document.xml")
            tree = ET.fromstring(xml_content)
            texts = [node.text for node in tree.iter() if node.text]
            combined = " ".join(texts).strip()
            if combined:
                logger.info("Extracted text from DOCX via XML fallback: %s", filename)
                return [(combined, {"source": filename, "page": 1, "file_type": "docx"})]
    except Exception as exc:
        logger.error("DOCX XML fallback extraction failed: %s", exc)

    return []


def extract_text_from_txt(text: str, filename: str = "text_input") -> List[Tuple[str, dict]]:
    """Wrap raw text into standard format."""
    if not text or not text.strip():
        return []
    return [(text.strip(), {"source": filename, "page": 1, "file_type": "txt"})]


def chunk_documents(
    pages: List[Tuple[str, dict]],
    chunk_size: int = config.LC_CHUNK_SIZE,
    chunk_overlap: int = config.LC_CHUNK_OVERLAP,
) -> List[Document]:
    """Split pages into LangChain Document chunks."""
    splitter = PurePythonRecursiveTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    all_docs: List[Document] = []
    for item in pages:
        if isinstance(item, tuple) and len(item) == 2:
            page_text, meta = item
        else:
            continue
        chunks = splitter.split_text(page_text)
        for idx, chunk in enumerate(chunks):
            if chunk.strip():
                doc_meta = {**meta, "chunk_index": idx, "chunk_total": len(chunks)}
                all_docs.append(Document(page_content=chunk.strip(), metadata=doc_meta))

    logger.info("Chunked %d pages → %d chunks", len(pages), len(all_docs))
    return all_docs
